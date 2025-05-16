from docx import Document

doc = Document()
doc.add_heading('Student Information Form', 0)
doc.add_paragraph('Full Name: CaoHuyThinh')
doc.add_paragraph('Student ID: 20230069')
doc.add_paragraph('Major: CS')

doc.save('test_data/form_sample.docx')